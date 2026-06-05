"""
Rapporten Debutade - Web Applicatie
===================================

Dashboard met gecombineerde rapportage van kasboek en bankrekening.

Versie: 1.0
Datum: 2026-02-21
"""

from __future__ import annotations

from datetime import datetime, date
import getpass
import json
import logging
import os
import sys
import time
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from flask import Flask, jsonify, redirect, render_template, request, g
from openpyxl import load_workbook


if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except AttributeError:
        import io

        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = os.getenv(
    "DEBUTADE_CONFIG",
    os.path.abspath(os.path.join(SCRIPT_DIR, "..", "config.json")),
)


def load_config(config_path: str) -> dict[str, Any]:
    if not os.path.exists(config_path):
        raise FileNotFoundError(f"Configuratiebestand niet gevonden: {config_path}")

    with open(config_path, "r", encoding="utf-8") as config_file:
        root_config = json.load(config_file)

    shared = root_config.get("shared", {})
    rapporten = root_config.get("rapporten", {})
    bank = root_config.get("bankrekening", {})
    kas = root_config.get("kasboek", {})

    grootboek_directory = shared.get("grootboek_directory", "")

    bank_excel_file_name = (
        rapporten.get("bank_excel_file_name")
        or bank.get("excel_file_name")
        or shared.get("bank_excel_file_name")
        or ""
    )
    kas_excel_file_name = (
        rapporten.get("kas_excel_file_name")
        or kas.get("excel_file_name")
        or ""
    )

    bank_sheets = rapporten.get("bank_sheets") or bank.get("required_sheets") or [
        "Bankrekening",
        "Spaarrekening 1",
        "Spaarrekening 2",
    ]
    kas_sheet_name = rapporten.get("kas_sheet_name") or kas.get("excel_sheet_name") or "Kas"

    def build_path(file_name: str) -> str:
        if not file_name:
            return ""
        if os.path.isabs(file_name):
            return file_name
        if grootboek_directory:
            return os.path.join(grootboek_directory, file_name)
        return file_name

    return {
        "bank_excel_path": build_path(bank_excel_file_name),
        "kas_excel_path": build_path(kas_excel_file_name),
        "grootboek_directory": grootboek_directory,
        "bank_sheets": bank_sheets,
        "kas_sheet_name": kas_sheet_name,
        "log_directory": shared.get("log_directory", os.path.join(SCRIPT_DIR, "logs")),
        "log_level": shared.get("log_level", "INFO"),
        "main_app_url": os.getenv("MAIN_APP_URL", "").strip(),
    }


try:
    config = load_config(CONFIG_PATH)
except (FileNotFoundError, KeyError, json.JSONDecodeError) as exc:
    print(f"WAARSCHUWING: {exc}")
    config = {
        "bank_excel_path": "",
        "kas_excel_path": "",
        "grootboek_directory": "",
        "bank_sheets": ["Bankrekening", "Spaarrekening 1", "Spaarrekening 2"],
        "kas_sheet_name": "Kas",
        "log_directory": os.path.join(SCRIPT_DIR, "logs"),
        "log_level": "INFO",
        "main_app_url": os.getenv("MAIN_APP_URL", "").strip(),
    }


LOG_DIRECTORY = config["log_directory"]
LOG_LEVEL = config["log_level"]
MAIN_APP_URL = config["main_app_url"]

app = Flask(
    __name__,
    template_folder=os.path.join(SCRIPT_DIR, "templates"),
    static_folder=os.path.join(SCRIPT_DIR, "static"),
)
app.config["TEMPLATES_AUTO_RELOAD"] = True

CACHE_TTL_SECONDS = int(os.getenv("DEBUTADE_CACHE_TTL_SECONDS", "20"))
BENCHMARK_ENABLED = os.getenv("DEBUTADE_BENCHMARK", "1") == "1"
RUNTIME_CACHE: dict[tuple[Any, ...], dict[str, Any]] = {}


def _file_signature(file_path: str) -> tuple[int, int] | None:
    if not file_path or not os.path.exists(file_path):
        return None
    stat = os.stat(file_path)
    return (stat.st_mtime_ns, stat.st_size)


def _cache_get(key: tuple[Any, ...]) -> Any:
    entry = RUNTIME_CACHE.get(key)
    if not entry:
        return None

    if time.time() - entry["ts"] > CACHE_TTL_SECONDS:
        RUNTIME_CACHE.pop(key, None)
        return None

    return entry["value"]


def _cache_set(key: tuple[Any, ...], value: Any) -> Any:
    RUNTIME_CACHE[key] = {"ts": time.time(), "value": value}
    return value


def normalize_text(value: Any) -> str:
    return str(value or "").strip()


def parse_date(value: Any) -> datetime | None:
    if isinstance(value, datetime):
        return value
    if isinstance(value, date):
        return datetime(value.year, value.month, value.day)
    if value is None:
        return None

    raw = normalize_text(value)
    if not raw:
        return None

    for fmt in ("%d-%m-%Y", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%y", "%Y/%m/%d"):
        try:
            return datetime.strptime(raw, fmt)
        except ValueError:
            continue
    return None


def parse_amount(value: Any) -> float | None:
    if isinstance(value, (int, float)):
        return float(value)

    raw = normalize_text(value)
    if not raw:
        return None

    raw = raw.replace("€", "").replace(" ", "")
    if raw.count(",") == 1 and raw.count(".") > 1:
        raw = raw.replace(".", "")
    raw = raw.replace(".", "").replace(",", ".")

    try:
        return float(raw)
    except ValueError:
        return None


def source_from_sheet(sheet_name: str, is_kas: bool = False) -> str:
    if is_kas:
        return "Kas"

    lower_sheet = normalize_text(sheet_name).lower()
    if "spaarrekening 1" in lower_sheet:
        return "Spaarrekening 1"
    if "spaarrekening 2" in lower_sheet:
        return "Spaarrekening 2"
    if "bank" in lower_sheet:
        return "Bankrekening"
    return normalize_text(sheet_name) or "Onbekend"


def read_transactions_from_sheet(file_path: str, sheet_name: str, is_kas: bool = False) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    workbook = None

    if not file_path or not os.path.exists(file_path):
        return records

    try:
        workbook = load_workbook(file_path, read_only=True, data_only=True)
        if sheet_name not in workbook.sheetnames:
            return records

        sheet = workbook[sheet_name]
        source = source_from_sheet(sheet_name, is_kas=is_kas)

        header_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
        header_map: dict[str, int] = {}
        if header_row:
            for index, cell_value in enumerate(header_row):
                key = normalize_text(cell_value).lower()
                if key and key not in header_map:
                    header_map[key] = index

        def get_value(row: tuple[Any, ...], header_names: tuple[str, ...], fallback_index: int | None = None) -> Any:
            for header_name in header_names:
                idx = header_map.get(header_name.lower())
                if idx is not None and idx < len(row):
                    return row[idx]
            if fallback_index is not None and fallback_index < len(row):
                return row[fallback_index]
            return None

        for row in sheet.iter_rows(min_row=2, values_only=True):
            if not row:
                continue

            txn_date = parse_date(get_value(row, ("datum",), 0))
            af_bij = normalize_text(get_value(row, ("af bij", "af/bij"), 5))
            amount = parse_amount(get_value(row, ("bedrag (eur)", "bedrag", "amount"), 6))

            if amount is None or af_bij not in {"Af", "Bij"}:
                continue

            sign = -1.0 if af_bij == "Af" else 1.0
            signed_amount = round(sign * amount, 2)
            month_key = txn_date.strftime("%Y-%m") if txn_date else "Onbekend"

            records.append(
                {
                    "date": txn_date.strftime("%Y-%m-%d") if txn_date else "",
                    "month": month_key,
                    "description": normalize_text(get_value(row, ("naam / omschrijving", "omschrijving", "naam"), 1)),
                    "source": source,
                    "af_bij": af_bij,
                    "amount": round(float(amount), 2),
                    "signed_amount": signed_amount,
                    "bon": normalize_text(get_value(row, ("bon",), 10)),
                    "tag": normalize_text(get_value(row, ("tag",), 11)) or "(Geen tag)",
                    "mededelingen": normalize_text(get_value(row, ("mededelingen",), 8)),
                }
            )

        return records
    except Exception as exc:
        logging.error("Fout bij lezen sheet %s uit %s: %s", sheet_name, file_path, exc)
        return records
    finally:
        if workbook:
            workbook.close()


def load_all_transactions() -> tuple[list[dict[str, Any]], list[str]]:
    bank_file = config.get("bank_excel_path", "")
    kas_file = config.get("kas_excel_path", "")
    bank_sheets = tuple(config.get("bank_sheets", []))
    kas_sheet_name = config.get("kas_sheet_name", "Kas")

    cache_key = (
        "all_transactions",
        bank_file,
        _file_signature(bank_file),
        kas_file,
        _file_signature(kas_file),
        bank_sheets,
        kas_sheet_name,
    )
    cached = _cache_get(cache_key)
    if cached is not None:
        return cached

    all_records: list[dict[str, Any]] = []
    warnings: list[str] = []

    if not bank_file or not os.path.exists(bank_file):
        warnings.append("Bank Excel bestand niet gevonden of niet ingesteld.")
    if not kas_file or not os.path.exists(kas_file):
        warnings.append("Kas Excel bestand niet gevonden of niet ingesteld.")

    for sheet_name in bank_sheets:
        all_records.extend(read_transactions_from_sheet(bank_file, sheet_name, is_kas=False))

    all_records.extend(read_transactions_from_sheet(kas_file, kas_sheet_name, is_kas=True))

    all_records.sort(key=lambda item: (item["date"], item["description"]), reverse=True)
    return _cache_set(cache_key, (all_records, warnings))


def get_report_payload_cached() -> dict[str, Any]:
    records, warnings = load_all_transactions()
    bank_file = config.get("bank_excel_path", "")
    kas_file = config.get("kas_excel_path", "")

    cache_key = (
        "report_payload",
        bank_file,
        _file_signature(bank_file),
        kas_file,
        _file_signature(kas_file),
    )
    cached = _cache_get(cache_key)
    if cached is not None:
        return cached

    months = sorted({row["month"] for row in records if row["month"] and row["month"] != "Onbekend"})
    tags = sorted({row["tag"] for row in records if row["tag"]})
    sources = sorted({row["source"] for row in records if row["source"]})

    payload = {
        "transactions": records,
        "filters": {
            "months": months,
            "tags": tags,
            "sources": sources,
        },
        "warnings": warnings,
    }
    return _cache_set(cache_key, payload)


def get_export_directory() -> str:
    export_dir = normalize_text(config.get("grootboek_directory", ""))
    return export_dir or SCRIPT_DIR


def _set_docx_cell_text(cell, value: Any, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(value or ""))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_docx_hyperlink(document, paragraph, url: str, text: str) -> None:
    """Add a hyperlink to a paragraph in a Word document with proper relationship."""
    from docx.oxml import parse_xml
    from docx.shared import RGBColor
    
    if not url or not text:
        run = paragraph.add_run(text or url or "")
        run.font.size = Pt(9)
        return
    
    # Add relationship to document for the hyperlink
    try:
        rel_id = document.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    except:
        # Fallback if relationship fails
        run = paragraph.add_run(text)
        run.font.size = Pt(9)
        return
    
    # Create hyperlink XML with the relationship ID
    hyperlink_xml = (
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        f'r:id="{rel_id}">'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rStyle w:val="Hyperlink"/>'
        f'<w:u w:val="single"/>'
        f'<w:color w:val="0563C1"/>'
        f'</w:rPr>'
        f'<w:t>{text}</w:t>'
        f'</w:r>'
        f'</w:hyperlink>'
    )
    
    try:
        hyperlink_element = parse_xml(hyperlink_xml)
        paragraph._element.append(hyperlink_element)
        
        # Format the text
        for run in paragraph.runs:
            if text in run.text or run.text == text:
                run.font.size = Pt(9)
                run.underline = True
                run.font.color.rgb = RGBColor(5, 99, 193)  # Blue color
                break
    except Exception as e:
        # Fallback: just add plain text if hyperlink fails
        run = paragraph.add_run(text)
        run.font.size = Pt(9)


def _set_repeat_table_header(table) -> None:
    """Set the first row of a table to repeat on each page."""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Add repeat header rows property
    tblPrChange = OxmlElement('w:tblHeader')
    tblPr.append(tblPrChange)
    
    # Set first row as header
    if len(table.rows) > 0:
        tr = table.rows[0]._element
        trPr = tr.trPr
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        
        repeat_header = OxmlElement('w:tblHeader')
        trPr.append(repeat_header)


def write_transactions_docx(file_path: str, payload: dict[str, Any]) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(9)

    document.add_heading("Rapportage Kasboek & Bankrekening - Tabelexport", level=1)

    generated_at = normalize_text(payload.get("generatedAt")) or datetime.now().strftime("%d-%m-%Y %H:%M:%S")
    total_transactions = int(payload.get("totalTransactions") or 0)
    filters = payload.get("filters") if isinstance(payload.get("filters"), dict) else {}

    meta_paragraph = document.add_paragraph()
    meta_paragraph.add_run(f"Gegenereerd op {generated_at} | Totaal transacties: {total_transactions}").bold = True

    filter_lines = [
        f"Maand: {normalize_text(filters.get('month')) or 'Alle maanden'}",
        f"Tag filter: {normalize_text(filters.get('tag')) or 'Alle tags'}",
        f"Bron filter: {normalize_text(filters.get('source')) or 'Alle bronnen'}",
        f"Zoeken mededelingen: {normalize_text(filters.get('mededelingen')) or '-'}",
    ]
    for line in filter_lines:
        document.add_paragraph(line)

    document.add_paragraph("")

    # Set column widths for portrait format
    column_widths = [Cm(1.4), Cm(1.2), Cm(1.2), Cm(1.2), Cm(2.8), Cm(2.0), Cm(0.9)]
    headers = ["Datum", "Bron", "Debet", "Credit", "Omschrijving", "Mededelingen", "Bon"]

    groups = payload.get("groups") if isinstance(payload.get("groups"), list) else []
    if not groups:
        # Create a single table for the "no transactions" message
        table = document.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, width in enumerate(column_widths):
            for row in table.rows:
                row.cells[idx].width = width
        
        empty_cells = table.rows[0].cells
        merged_cell = empty_cells[0].merge(empty_cells[6])
        _set_docx_cell_text(merged_cell, "Geen transacties gevonden.")
    else:
        # Create a separate table for each category/tag
        for group in groups:
            tag = normalize_text(group.get("tag")) or "(Geen tag)"
            rows = group.get("rows") if isinstance(group.get("rows"), list) else []

            # Add category heading
            category_heading = document.add_paragraph()
            category_heading.add_run(f"Categorie: {tag} ({len(rows)} transacties)").bold = True

            # Create table for this category
            table = document.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Set column widths
            for idx, width in enumerate(column_widths):
                for row in table.rows:
                    row.cells[idx].width = width

            # Add headers
            header_cells = table.rows[0].cells
            for index, header in enumerate(headers):
                _set_docx_cell_text(header_cells[index], header, bold=True)

            # Set header to repeat on each page
            _set_repeat_table_header(table)

            # Add transaction rows for this category
            for row in rows:
                row_cells = table.add_row().cells
                _set_docx_cell_text(row_cells[0], row.get("date"))
                _set_docx_cell_text(row_cells[1], row.get("source"))
                _set_docx_cell_text(row_cells[2], row.get("debet"))
                _set_docx_cell_text(row_cells[3], row.get("credit"))
                _set_docx_cell_text(row_cells[4], row.get("description"))
                _set_docx_cell_text(row_cells[5], row.get("mededelingen"))
                
                # Handle bon column - only show as link if bon URL is available
                bon_url = normalize_text(row.get("bon"))
                bon_cell = row_cells[6]
                bon_cell.text = ""
                bon_paragraph = bon_cell.paragraphs[0]
                bon_paragraph.alignment = 1  # Center alignment
                
                if bon_url and bon_url.startswith("http"):
                    _add_docx_hyperlink(document, bon_paragraph, bon_url, "BON")
                else:
                    # Leave empty if no bon URL
                    _set_docx_cell_text(bon_cell, "")
                
                bon_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Add subtotal row
            subtotal_cells = table.add_row().cells
            subtotal_label_cell = subtotal_cells[0].merge(subtotal_cells[1])
            _set_docx_cell_text(subtotal_label_cell, f"Subtotaal {tag}", bold=True)
            _set_docx_cell_text(subtotal_cells[2], group.get("subtotalDebet"), bold=True)
            _set_docx_cell_text(subtotal_cells[3], group.get("subtotalCredit"), bold=True)
            _set_docx_cell_text(subtotal_cells[4], "")
            _set_docx_cell_text(subtotal_cells[5], "")
            _set_docx_cell_text(subtotal_cells[6], "")

            # Add spacing between tables
            document.add_paragraph("")

    document.save(file_path)


@app.before_request
def log_request() -> None:
    logging.info("REQUEST %s %s %s", request.remote_addr, request.method, request.path)
    if BENCHMARK_ENABLED:
        g._start_time = time.perf_counter()


@app.after_request
def benchmark_request(response):
    if BENCHMARK_ENABLED and hasattr(g, "_start_time"):
        elapsed_ms = (time.perf_counter() - g._start_time) * 1000
        logging.info("PERF %s %s %s %.1fms", request.method, request.path, response.status_code, elapsed_ms)
    return response


@app.route("/")
def index():
    return render_template(
        "index.html",
        current_date=datetime.now().strftime("%d-%m-%Y"),
        current_user=getpass.getuser(),
        main_app_url=MAIN_APP_URL,
    )


@app.route("/api/report-data")
def report_data():
    return jsonify(get_report_payload_cached())


@app.route("/api/export-table-docx", methods=["POST"])
def export_table_docx():
    payload = request.get_json(silent=True) or {}
    groups = payload.get("groups")

    if not isinstance(groups, list):
        return jsonify({"success": False, "message": "Ongeldige exportdata ontvangen."}), 400

    export_dir = get_export_directory()

    try:
        os.makedirs(export_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"transacties debutade {timestamp}.docx"
        file_path = os.path.join(export_dir, file_name)

        write_transactions_docx(file_path, payload)
        logging.info("DOCX export aangemaakt: %s", file_path)
        return jsonify(
            {
                "success": True,
                "message": "Word-document aangemaakt.",
                "file_name": file_name,
                "file_path": file_path,
            }
        )
    except Exception as exc:
        logging.exception("Fout bij DOCX export: %s", exc)
        return jsonify({"success": False, "message": f"DOCX export mislukt: {exc}"}), 500


@app.route("/quit", methods=["POST"])
def quit_application():
    try:
        user = getpass.getuser()
        logging.info("APPLICATIE AFGESLOTEN | Gebruiker: %s", user)
        logging.info("=" * 70)

        response = jsonify({"success": True, "message": "Applicatie sluit af"})

        def shutdown_server() -> None:
            import time

            time.sleep(1)
            logging.info("Flask server wordt beeindigd...")
            os._exit(0)

        import threading

        shutdown_thread = threading.Thread(target=shutdown_server, daemon=True)
        shutdown_thread.start()

        return response, 200
    except Exception as exc:
        logging.error("Fout bij afsluiten applicatie: %s", str(exc))
        return jsonify({"success": False, "message": f"Fout: {str(exc)}"}), 500


@app.route("/settings")
def settings():
    if MAIN_APP_URL:
        return redirect(f"{MAIN_APP_URL}/settings")
    return jsonify({"success": False, "message": "Instellingen zijn alleen beschikbaar via de hoofdapp."}), 403


if __name__ == "__main__":
    if not os.path.exists(LOG_DIRECTORY):
        try:
            os.makedirs(LOG_DIRECTORY)
        except Exception as exc:
            print(f"FOUT: Kan log directory niet aanmaken: {LOG_DIRECTORY}")
            print(f"Details: {str(exc)}")
            exit(1)

    log_file_path = os.path.join(LOG_DIRECTORY, "rapporten_webapp_log.txt")
    logging.basicConfig(
        filename=log_file_path,
        level=getattr(logging, str(LOG_LEVEL).upper(), logging.INFO),
        format="%(asctime)s - %(levelname)s - %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )

    logging.info("=" * 70)
    logging.info("RAPPORTEN START")
    logging.info("Bank bestand: %s", config.get("bank_excel_path", ""))
    logging.info("Kas bestand: %s", config.get("kas_excel_path", ""))
    logging.info("=" * 70)

    port = int(os.getenv("DEBUTADE_APP_PORT", "5004"))
    app.run(debug=False, host="127.0.0.1", port=port, use_reloader=False)
